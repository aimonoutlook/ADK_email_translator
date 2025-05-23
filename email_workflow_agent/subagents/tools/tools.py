# email-agent-workflow/email_workflow_agent/tools/tools.py
import os
import uuid
import logging
import requests # Example for external API calls
from typing import Any, Dict, Optional, List
from google.adk.tools import FunctionTool, ToolContext
from google.genai import types
# Import the sensitive data handling callbacks
from .callbacks import handle_sensitive_before, handle_sensitive_after

logger = logging.getLogger(__name__)

# --- Custom Tool Functions ---

# Tool 1: Download and Save Attachments as Artifacts
# This tool is called by the Custom Orchestrator agent
async def download_attachments(tool_context: ToolContext) -> Dict[str, Any]:
    """
    Tool to download attachments from the initial email state
    and save them as artifacts.

    Reads from state['initial_attachments'].
    Writes artifact filenames/versions to state['attachment_artifacts'].
    Writes original file format to state['original_file_format'] for later use.
    """
    logger.info(f"[Tool] download_attachments called.")
    initial_attachments = tool_context.state.get("initial_attachments", [])
    saved_artifact_details = {}
    original_file_format = None # Assuming first attachment dictates format

    if not initial_attachments:
        logger.warning(f"[Tool] No attachments found in initial state.")
        tool_context.state["attachment_artifacts"] = saved_artifact_details # Save empty dict
        tool_context.state["original_file_format"] = original_file_format
        return {"status": "success", "message": "No attachments to process."}

    # In a real scenario, `initial_attachments` would contain file data (bytes)
    # or temporary paths accessible to the agent.
    # For this demo, we simulate by creating dummy artifacts.
    logger.info(f"[Tool] Simulating download and saving {len(initial_attachments)} attachments as artifacts.")

    for filename in initial_attachments:
        # Simulate reading file bytes (replace with actual file read/download logic)
        dummy_bytes = b"Dummy content for " + filename.encode('utf-8')
        # Determine mime type (replace with actual detection or mapping)
        mime_type = "application/octet-stream"
        if filename.lower().endswith(".docx"):
             mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
             if original_file_format is None: original_file_format = "docx"
        elif filename.lower().endswith(".pdf"):
             mime_type = "application/pdf"
             if original_file_format is None: original_file_format = "pdf"
        # Add other types as needed

        artifact_part = types.Part.from_data(data=dummy_bytes, mime_type=mime_type)

        try:
            # Save the artifact using the context method
            # Use original filename, ArtifactService handles versioning and scoping
            version = await tool_context.save_artifact(filename=filename, artifact=artifact_part)
            saved_artifact_details[filename] = version
            logger.info(f"[Tool] Saved artifact '{filename}' version {version}.")
        except ValueError as e:
             logger.error(f"[Tool] Error saving artifact '{filename}': {e}")
             return {"status": "error", "message": f"Failed to save artifact {filename}: {e}"}
        except Exception as e:
             logger.error(f"[Tool] Unexpected error saving artifact '{filename}': {e}")
             return {"status": "error", "message": f"Failed to save artifact {filename}: {e}"}


    # Update state with the names and versions of the saved artifacts
    tool_context.state["attachment_artifacts"] = saved_artifact_details
    # Store the format of the first attachment for later conversion/editing
    tool_context.state["original_file_format"] = original_file_format

    return {"status": "success", "message": f"Attachments processed and saved as artifacts.", "artifacts": saved_artifact_details}

# Wrap the tool function in a FunctionTool
download_attachments_tool = FunctionTool(func=download_attachments)


# Tool 2: Extract Text from Document Artifacts
# This tool is called by the Custom Orchestrator agent
async def extract_text(tool_context: ToolContext) -> Dict[str, Any]:
    """
    Tool to extract text from document artifacts.

    Reads artifact names/versions from state['attachment_artifacts'].
    Loads artifacts, extracts text (handles .docx, .pdf etc.).
    Writes extracted text(s) and potentially original format to state.
    """
    logger.info(f"[Tool] extract_text called.")
    attachment_artifacts = tool_context.state.get("attachment_artifacts", {})
    extracted_texts = {} # Use a dict to store text from multiple files if needed
    original_file_format = None

    if not attachment_artifacts:
        logger.warning(f"[Tool] No attachment artifacts found in state.")
        tool_context.state["extracted_text"] = "" # Save empty string
        return {"status": "success", "message": "No artifacts to extract text from."}

    # Assuming we only process the first document attachment for simplicity in this workflow
    # In a real app, you might need to process multiple documents (e.g., original and translated)
    first_artifact_name = list(attachment_artifacts.keys())[0]
    first_artifact_version = attachment_artifacts[first_artifact_name]

    try:
        # Load the artifact content
        artifact_part = await tool_context.load_artifact(
            filename=first_artifact_name, version=first_artifact_version
        )

        if not artifact_part or not artifact_part.inline_data:
             logger.error(f"[Tool] Failed to load or artifact has no inline data: {first_artifact_name} v{first_artifact_version}.")
             return {"status": "error", "message": f"Failed to load artifact {first_artifact_name} for text extraction."}

        # Extract text based on MIME type
        mime_type = artifact_part.inline_data.mime_type
        file_content_bytes = artifact_part.inline_data.data
        extracted_text_content = ""

        # --- Placeholder: Text Extraction Logic ---
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            logger.info(f"[Tool] Extracting text from DOCX: {first_artifact_name}")
            try:
                from docx import Document # Requires python-docx
                from io import BytesIO
                doc = Document(BytesIO(file_content_bytes))
                extracted_text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                original_file_format = "docx"
                logger.info(f"[Tool] Extracted {len(extracted_text_content)} chars from DOCX.")
            except ImportError:
                logger.error("[Tool] python-docx not installed. Cannot extract text from DOCX.")
                return {"status": "error", "message": "Python-docx library not found. Cannot extract text from DOCX."}
            except Exception as e:
                logger.error(f"[Tool] Error extracting text from DOCX: {e}")
                return {"status": "error", "message": f"Error extracting text from DOCX: {e}"}

        elif mime_type == "application/pdf":
             logger.info(f"[Tool] Extracting text from PDF: {first_artifact_name}")
             try:
                 from PyPDF2 import PdfReader # Requires PyPDF2
                 from io import BytesIO
                 reader = PdfReader(BytesIO(file_content_bytes))
                 extracted_text_content = "".join([page.extract_text() for page in reader.pages if page.extract_text()])
                 original_file_format = "pdf"
                 logger.info(f"[Tool] Extracted {len(extracted_text_content)} chars from PDF.")
             except ImportError:
                 logger.error("[Tool] PyPDF2 not installed. Cannot extract text from PDF.")
                 return {"status": "error", "message": "PyPDF2 library not found. Cannot extract text from PDF."}
             except Exception as e:
                 logger.error(f"[Tool] Error extracting text from PDF: {e}")
                 return {"status": "error", "message": f"Error extracting text from PDF: {e}"}
        else:
            logger.warning(f"[Tool] Unsupported MIME type for text extraction: {mime_type} for {first_artifact_name}. Attempting raw text.")
            try:
                 # Attempt decoding as text if possible (e.g., .txt files or simple encodings)
                 extracted_text_content = file_content_bytes.decode('utf-8', errors='ignore')
                 original_file_format = "txt" # Assume if decode works
                 logger.info(f"[Tool] Decoded {len(extracted_text_content)} chars as text.")
            except Exception as e:
                 logger.error(f"[Tool] Failed to decode content as text: {e}")
                 return {"status": "error", "message": f"Unsupported file type for text extraction: {mime_artifact.mime_type}"}
        # --- End Placeholder ---

        # Save the extracted text (and original format) to state
        tool_context.state["extracted_text"] = extracted_text_content
        tool_context.state["original_file_format"] = original_file_format

        return {"status": "success", "message": f"Text extracted from {first_artifact_name}.", "extracted_char_count": len(extracted_text_content)}

    except Exception as e:
        logger.error(f"[Tool] Unexpected error during text extraction: {e}")
        return {"status": "error", "message": f"Unexpected error during text extraction: {e}"}

# Wrap the tool function
extract_text_tool = FunctionTool(func=extract_text)


# Tool 3: Translate Text (Applies Sensitive Data Callbacks)
# Called by TranslationWorkflowAgent
# Attach the sensitive data callbacks to this tool
async def translate_text(tool_context: ToolContext, text: str, target_language: str = "French") -> Dict[str, Any]:
    """
    Tool to translate text using an external API.
    Sensitive data handling callbacks are attached to this tool.

    Reads text argument (may contain placeholders).
    Reads target_language argument.
    Returns translated text (may contain placeholders initially).
    """
    logger.info(f"[Tool] translate_text called for {len(text)} chars to {target_language}.")

    # --- Placeholder: Call External Translation API ---
    try:
        # In a real app, you'd call an API like Google Cloud Translation, DeepL, etc.
        # Example using requests (install it first: pip install requests)
        # api_url = "https://translation.googleapis.com/language/translate/v2" # Example
        # api_key = os.getenv("TRANSLATION_API_KEY") # Get from .env
        # if not api_key:
        #     logger.error("[Tool] TRANSLATION_API_KEY not set.")
        #     return {"status": "error", "message": "Translation API key not configured."}
        #
        # response = requests.post(api_url, params={'key': api_key}, json={
        #     'q': text,
        #     'target': target_language,
        #     'source': 'en', # Assuming source is always English based on prompt
        # })
        # response.raise_for_status() # Raise an exception for bad status codes
        # result = response.json()
        # translated_text_content = result['data']['translations'][0]['translatedText']

        # Simulate translation (placeholders are carried through)
        translated_text_content = f"Translated: {text} (to {target_language})" # Placeholders like __VAR_1__ will be here

        logger.info(f"[Tool] Simulated translation.")
        return {"status": "success", "translated_text": translated_text_content}

    except Exception as e:
        logger.error(f"[Tool] Error calling translation API: {e}")
        return {"status": "error", "message": f"Translation failed: {e}"}
    # --- End Placeholder ---

# Wrap the tool function and ATTACH THE SENSITIVE DATA CALLBACKS
translate_text_tool = FunctionTool(
    func=translate_text,
    before_tool_callback=handle_sensitive_before, # Apply sensitive data handler before
    after_tool_callback=handle_sensitive_after,   # Apply sensitive data handler after
)


# Tool 4: Check Translation Quality (Applies Sensitive Data Callbacks)
# Called by TranslationWorkflowAgent or ReviewWorkflowAgent
# Attach the sensitive data callbacks to this tool
async def check_translation(tool_context: ToolContext, original_text: str, translated_text: str) -> Dict[str, Any]:
    """
    Tool to check the quality and accuracy of translated text against the original.
    Sensitive data handling callbacks are attached to this tool.

    Reads original_text argument (may contain placeholders).
    Reads translated_text argument (may contain placeholders).
    Returns feedback/score (may contain placeholders initially).
    """
    logger.info(f"[Tool] check_translation called for {len(original_text)} vs {len(translated_text)} chars.")

    # --- Placeholder: Translation Quality Check Logic ---
    try:
        # In a real app, you might use:
        # - Another LLM call specifically for evaluation
        # - NLP libraries for similarity scores
        # - Comparison against a Translation Memory (TM)
        # The texts passed here will have sensitive data replaced by placeholders *if* the callback ran.

        # Simulate quality check feedback (placeholders are carried through)
        feedback_text = f"Quality check feedback: The translation seems reasonable, but pay attention to terms like {translated_text.split(' ')[-1]} (from original {original_text.split(' ')[-1]}). Score: 85/100."
        # Placeholders like __VAR_1__ will be in original_text/translated_text args
        # and potentially generated in feedback_text by the simulation here.

        logger.info(f"[Tool] Simulated quality check.")
        return {"status": "success", "feedback_text": feedback_text, "score": 85} # Example structure

    except Exception as e:
        logger.error(f"[Tool] Error during translation quality check: {e}")
        return {"status": "error", "message": f"Quality check failed: {e}"}
    # --- End Placeholder ---

# Wrap the tool function and ATTACH THE SENSITIVE DATA CALLBACKS
check_translation_tool = FunctionTool(
    func=check_translation,
    before_tool_callback=handle_sensitive_before, # Apply sensitive data handler before
    after_tool_callback=handle_sensitive_after,   # Apply sensitive data handler after
)


# Tool 5: Edit Word Document with Track Changes (Applies Sensitive Data Callbacks)
# Called by ReviewWorkflowAgent
# Attach the sensitive data callbacks to this tool
async def edit_word_doc(tool_context: ToolContext, artifact_name: str, artifact_version: int, edit_instructions: str) -> Dict[str, Any]:
    """
    Tool to load a Word document artifact, apply edits with track changes,
    and save the edited document as a new artifact version.
    Sensitive data handling callbacks are attached to this tool.

    Reads document artifact by name/version.
    Reads edit_instructions argument (may contain placeholders).
    Saves edited document as a new artifact version.
    Returns new artifact name/version.
    """
    logger.info(f"[Tool] edit_word_doc called for artifact '{artifact_name}' v{artifact_version}.")

    try:
        # Load the document artifact content
        doc_artifact_part = await tool_context.load_artifact(filename=artifact_name, version=artifact_version)

        if not doc_artifact_part or not doc_artifact_part.inline_data:
             logger.error(f"[Tool] Failed to load document artifact or it has no inline data: {artifact_name} v{artifact_version}.")
             return {"status": "error", "message": f"Failed to load document artifact {artifact_name} for editing."}

        # --- Placeholder: Document Editing Logic with Track Changes ---
        try:
             # This is complex! python-docx has limited track changes support.
             # You might need a commercial library or API (e.g., Aspose.Words, DocuSign, Google Docs API)
             # Or a workflow involving converting to a different format, editing, and converting back.

             logger.info(f"[Tool] Applying edits with track changes based on instructions ({len(edit_instructions)} chars).")
             edited_content_bytes = b"Edited document content with track changes simulated.\n" + doc_artifact_part.inline_data.data # Simulate edit
             edited_mime_type = doc_artifact_part.inline_data.mime_type # Keep original type

             # Note: Real implementation involves parsing docx XML, applying changes, saving.
             # Placeholder text may contain placeholders *if* the callback processed it.
             # If you need to process the *document content* itself for sensitive data within the tool,
             # you'd need to do it here before calling an external editing API, and apply callbacks
             # around that internal content processing step, not just the tool args/response.
             # For simplicity, we assume sensitive data is mainly in instructions and output text feedback.

        except ImportError:
             logger.error("[Tool] python-docx not installed. Cannot edit DOCX.")
             return {"status": "error", "message": "Python-docx library not found. Cannot edit DOCX."}
        except Exception as e:
             logger.error(f"[Tool] Error during document editing: {e}")
             return {"status": "error", "message": f"Document editing failed: {e}"}
        # --- End Placeholder ---

        # Create a new artifact part for the edited document
        edited_artifact_part = types.Part.from_data(data=edited_content_bytes, mime_type=edited_mime_type)

        # Save the edited document as a *new version* of the same artifact filename
        # The artifact service handles assigning the next version number
        new_version = await tool_context.save_artifact(filename=artifact_name, artifact=edited_artifact_part)
        logger.info(f"[Tool] Saved edited document as artifact '{artifact_name}' version {new_version}.")

        return {"status": "success", "message": f"Document edited and saved as version {new_version}.", "edited_artifact_name": artifact_name, "edited_artifact_version": new_version}

    except Exception as e:
        logger.error(f"[Tool] Unexpected error during document editing workflow: {e}")
        return {"status": "error", "message": f"Unexpected error during document editing workflow: {e}"}

# Wrap the tool function and ATTACH THE SENSITIVE DATA CALLBACKS
edit_word_doc_tool = FunctionTool(
    func=edit_word_doc,
    # Apply sensitive data handling before and after
    before_tool_callback=handle_sensitive_before,
    after_tool_callback=handle_sensitive_after,
)


# Tool 6: Convert Text to Word Document Artifact
# Called by TranslationWorkflowAgent
async def convert_to_word(tool_context: ToolContext, translated_text: str, original_format: str = "docx") -> Dict[str, Any]:
    """
    Tool to convert text into a Word document artifact.

    Reads translated_text argument.
    Reads original_format argument from state.
    Creates a Word document.
    Saves the document as an artifact.
    Returns new artifact name/version.
    """
    logger.info(f"[Tool] convert_to_word called for {len(translated_text)} chars, original format '{original_format}'.")

    # --- Placeholder: Convert Text to Word Logic ---
    try:
        # Requires python-docx or similar library
        if original_format != "docx":
            logger.warning(f"[Tool] Original format '{original_format}' not DOCX. Converting to DOCX anyway.")

        try:
            from docx import Document # Requires python-docx
            from io import BytesIO
            doc = Document()
            doc.add_paragraph(translated_text)
            # Add more complex formatting if needed based on original_format or template

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            word_bytes = buffer.getvalue()
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            logger.info(f"[Tool] Created DOCX document ({len(word_bytes)} bytes).")
        except ImportError:
             logger.error("[Tool] python-docx not installed. Cannot create DOCX.")
             return {"status": "error", "message": "Python-docx library not found. Cannot create DOCX."}
        except Exception as e:
             logger.error(f"[Tool] Error creating DOCX: {e}")
             return {"status": "error", "message": f"Error creating DOCX: {e}"}
        # --- End Placeholder ---

        # Create a new artifact part for the Word document
        word_artifact_part = types.Part.from_data(data=word_bytes, mime_type=mime_type)

        # Define a filename for the translated/edited document
        # Use a consistent naming convention, maybe based on original filename and status
        original_filename = tool_context.state.get("initial_attachments", [None])[0] # Get name of first attachment
        if original_filename:
             base_name = os.path.splitext(original_filename)[0]
             # Assuming this is for translated document:
             output_filename = f"{base_name}_translated.docx"
        else:
             output_filename = f"translated_document_{uuid.uuid4().hex[:6]}.docx"

        # Save the document as a new artifact
        # Versioning starts from 0 for this new filename
        version = await tool_context.save_artifact(filename=output_filename, artifact=word_artifact_part)
        logger.info(f"[Tool] Saved Word document as artifact '{output_filename}' version {version}.")

        return {"status": "success", "message": f"Document saved as artifact '{output_filename}' v{version}.", "artifact_name": output_filename, "artifact_version": version}

    except Exception as e:
        logger.error(f"[Tool] Unexpected error during Word conversion: {e}")
        return {"status": "error", "message": f"Unexpected error during Word conversion: {e}"}

# Wrap the tool function
convert_to_word_tool = FunctionTool(func=convert_to_word)


# Tool 7: Send Final Email
# Called by EmailSenderAgent
async def send_final_email(tool_context: ToolContext) -> Dict[str, Any]:
    """
    Tool to send the final email with the processed document attachment.

    Reads recipient email from state['email_sender_email'].
    Reads email body from state['initial_reply_text'].
    Reads final document artifact name/version from state
    (either 'translated_document_artifact' or 'edited_document_artifact').
    Loads artifact and sends email.
    """
    logger.info(f"[Tool] send_final_email called.")

    recipient_email = tool_context.state.get("email_sender_email")
    email_body_text = tool_context.state.get("initial_reply_text") # Use initial reply text
    email_subject = tool_context.state.get("email_subject") # Use original subject

    # Determine which artifact is the final document based on email type
    email_type = tool_context.state.get("email_type")
    if email_type == "translation":
         final_artifact_details = tool_context.state.get("translated_document_artifact")
    elif email_type == "review":
         final_artifact_details = tool_context.state.get("edited_document_artifact")
    else:
         logger.error(f"[Tool] Cannot send email, unknown email type: {email_type}")
         return {"status": "error", "message": "Cannot send email, unknown process type."}

    if not recipient_email or not email_body_text or not final_artifact_details:
         logger.error(f"[Tool] Cannot send email, missing recipient, body, or artifact details.")
         return {"status": "error", "message": "Cannot send email, missing required information."}

    final_artifact_name = final_artifact_details.get("artifact_name")
    final_artifact_version = final_artifact_details.get("artifact_version")

    if not final_artifact_name or final_artifact_version is None:
         logger.error(f"[Tool] Cannot send email, final artifact details incomplete.")
         return {"status": "error", "message": "Cannot send email, final artifact details incomplete."}


    try:
        # Load the final document artifact content
        final_doc_artifact_part = await tool_context.load_artifact(
            filename=final_artifact_name, version=final_artifact_version
        )

        if not final_doc_artifact_part or not final_doc_artifact_part.inline_data:
            logger.error(f"[Tool] Failed to load final document artifact: {final_artifact_name} v{final_artifact_version}.")
            return {"status": "error", "message": f"Failed to load final document artifact {final_artifact_name}."}

        # --- Placeholder: Send Email Logic ---
        try:
            logger.info(f"[Tool] Simulating sending email to {recipient_email}.")
            logger.info(f"[Tool] Subject: {email_subject}")
            logger.info(f"[Tool] Body: {email_body_text}")
            logger.info(f"[Tool] Attaching artifact: {final_artifact_name} ({final_doc_artifact_part.inline_data.mime_type})")

            # In a real app, use an email sending library or API (e.g., SendGrid, Mailgun, Gmail API)
            # You would attach final_doc_artifact_part.inline_data.data (bytes) with the specified mime_type and filename.

            # Example using print for simulation
            print(f"\n--- SIMULATING EMAIL SEND ---")
            print(f"To: {recipient_email}")
            print(f"Subject: {email_subject}")
            print(f"Body:\n{email_body_text}")
            print(f"Attachment: {final_artifact_name} ({len(final_doc_artifact_part.inline_data.data)} bytes, {final_doc_artifact_part.inline_data.mime_type})")
            print(f"-----------------------------\n")

            # Simulate success
            logger.info(f"[Tool] Email simulation successful.")
            return {"status": "success", "message": "Email sent successfully."}

        except Exception as e:
            logger.error(f"[Tool] Error sending email: {e}")
            return {"status": "error", "message": f"Email sending failed: {e}"}
        # --- End Placeholder ---

    except Exception as e:
        logger.error(f"[Tool] Unexpected error loading artifact or sending email: {e}")
        return {"status": "error", "message": f"Unexpected error during email sending: {e}"}


# Wrap the tool function
send_email_tool = FunctionTool(func=send_final_email)